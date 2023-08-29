import pandas as pd
from docx import Document

from vocab_tool_functions import import_and_clean_data, lookup_multiple_words_on_letonika
from vocab_tool_functions import new_word

pd.set_option('display.max_columns', None)


def remove_unnecessary_characters(row):
    return row['conjugations'].replace('(', '').replace(')', '').replace("'", "").replace(' ', '')


def split_conjugations(_df):
    split_conjugations_df = _df['conjugations'].str.split(',', expand=True)
    return split_conjugations_df


def refine_pamatforma_col(row):
    if pd.isnull(row['pamatforma']):
        return row['word']
    else:
        return row['pamatforma']


def refine_conjugations_col(row):
    if len(row['conjugations']) == 0:
        return row['pamatforma']
    else:
        return row['conjugations']


def min_of_two_cols(row):
    return min(row['index'], row['min_index'])


def find_missing_search_words(_first_word_loc_df, _sw_to_pf_df):
    df_all = _first_word_loc_df.merge(_sw_to_pf_df.drop_duplicates(), left_on='word', right_on='search_word',
                                      how='left', indicator=True)
    df_all = df_all[df_all['_merge'] == 'left_only']['word']
    return df_all


# create the list of words
df = import_and_clean_data('enter_text.txt').drop(columns='other')
df = df.reset_index()
first_word_loc_df = df[['word', 'index']].groupby('word').agg('min').reset_index()


def run_round_of_lookups():
    pf_to_conjs_df = pd.read_csv('pf_to_conjs.csv')
    sw_to_pf_df = pd.read_csv('sw_to_pf_df.csv')
    search_words = find_missing_search_words(first_word_loc_df, sw_to_pf_df)
    if len(search_words) == 1:
        print('Done searching')
        return 1

    else:
        search_words = search_words[:100]
        new_pf_to_conjs_df, new_sw_to_pf_df = lookup_multiple_words_on_letonika(search_words)
        pf_to_conjs_df = pf_to_conjs_df.append(new_pf_to_conjs_df, ignore_index=True).drop_duplicates().dropna()
        sw_to_pf_df = sw_to_pf_df.append(new_sw_to_pf_df, ignore_index=True).drop_duplicates().dropna()
        pf_to_conjs_df.to_csv('pf_to_conjs.csv', index=False)
        sw_to_pf_df.to_csv('sw_to_pf_df.csv', index=False)
        print('Round completed')
        return 0


x = 0
while x == 0:
    x = run_round_of_lookups()

pf_to_conjs_df = pd.read_csv('pf_to_conjs.csv')
sw_to_pf_df = pd.read_csv('sw_to_pf_df.csv')

unique_sw_to_pf_words = sw_to_pf_df.groupby('search_word').agg('count').reset_index()
unique_sw_to_pf_words = unique_sw_to_pf_words[unique_sw_to_pf_words['pamatforma'] == 1].drop(columns='pamatforma')
unique_sw_to_pf_words = unique_sw_to_pf_words.merge(sw_to_pf_df, on='search_word')
unique_sw_to_pf_words = unique_sw_to_pf_words.merge(first_word_loc_df, left_on='search_word', right_on='word') \
    .drop(columns='word')[['pamatforma', 'index']]
unique_sw_to_pf_words = unique_sw_to_pf_words.groupby('pamatforma').agg('min').reset_index()

word_indexes = sw_to_pf_df.merge(first_word_loc_df, left_on='search_word', right_on='word').drop(columns='word')
word_indexes.columns = ['search_word', 'pamatforma', 'first_search_word_index']
word_indexes = word_indexes.merge(unique_sw_to_pf_words, how='left')
word_indexes.columns = ['search_word', 'pamatforma', 'first_search_word_index', 'first_pamatforma_index']
word_indexes = word_indexes[['search_word', 'first_search_word_index', 'first_pamatforma_index']].fillna(-1). \
    groupby('search_word').agg('min').reset_index()


def return_min_index(row):
    if row['first_pamatforma_index'] == -1:
        return row['first_search_word_index']
    else:
        return min(row['first_search_word_index'], row['first_pamatforma_index'])


word_indexes['first_position'] = word_indexes.apply(return_min_index, axis=1)
word_indexes = word_indexes[['search_word', 'first_position']]
df = df.merge(word_indexes, left_on='word', right_on='search_word').sort_values(by='index').reset_index(drop=True). \
    drop(columns='search_word')
df.columns = ['position', 'word_original', 'word', 'first_position']
df['is_new_word'] = df.apply(new_word, axis=1)
df = df.reset_index(drop=True)

# di = {1: "red", 0: "green"}
# df['is_new_word'] = df['is_new_word'].map(di)
original_words = df['word_original']
colour_flag = df['is_new_word']

document = Document()
p = document.add_paragraph('')
for idx, word in enumerate(original_words):
    if colour_flag[idx] == 1:
        p.add_run(' ' + original_words[idx]).bold = True
    else:
        p.add_run(' ' + original_words[idx])

document.save('output.docx')

# create a viz
import matplotlib.pyplot as plt

df_viz = df[:]
df_viz['ma'] = df_viz[['is_new_word']].rolling(window=100).mean()
df_viz['space_between_new_words'] = 1 / df_viz['ma']
plt.plot(df_viz.space_between_new_words)
plt.show()

df_viz.tail(100)
