from bs4 import BeautifulSoup
import requests
import pandas as pd
import numpy as np
from collections import Counter
import matplotlib.pyplot as plt


pd.set_option('display.max_rows', 100000)
from termcolor import colored
from docx import Document
import re


def import_and_clean_data(file_name):
    file = open(file_name, encoding="utf-8-sig")
    words = file.read().split()
    df = pd.DataFrame(words)
    df.columns = ['word_original']
    df['word'] = df['word_original'].str.lower()
    df["word"] = df['word'].str.replace('[^\w\s]', '')
    df['other'] = 1
    return df


def create_agg_df(df):
    df_agg = df[['word', 'other']].groupby(['word']).agg(['sum'])
    df_agg.columns = ['word_count']
    df_agg = df_agg.sort_values(by='word_count', ascending=False)
    df_agg['cum_sum'] = df_agg['word_count'].cumsum()
    df_agg['cum_perc'] = 100 * df_agg['cum_sum'] / df_agg['word_count'].sum()
    df_agg = df_agg.reset_index()
    return df_agg


def create_study_df(df):
    study_df = df.reset_index()[['index', 'word']]
    study_df.columns = ['position', 'word']

    first_positions_df = study_df.groupby('word').agg('min').reset_index()
    first_positions_df.columns = ['word', 'first_position']
    first_positions_df = first_positions_df.sort_values('first_position').reset_index(drop=True)

    study_df = study_df.merge(first_positions_df).sort_values('position').reset_index(drop=True)

    study_df['new_word'] = study_df.apply(new_word, axis=1)
    study_df['vocab_size'] = study_df[['new_word']].cumsum()
    unique_word_df = study_df[study_df['new_word'] == 1][['word']].reset_index(drop=True)
    return study_df, unique_word_df


def new_word(row):
    if row['position'] == row['first_position']:
        return 1
    else:
        return 0


def output_bold_doc(file_name):
    df = import_and_clean_data(file_name).drop(columns='other')
    df = df.reset_index()
    first_word_loc_df = df[['word', 'index']].groupby('word').agg('min').reset_index()
    first_word_loc_df.columns = ['word', 'first_loc']
    df = df.merge(first_word_loc_df, on='word').sort_values(by='index')
    df.columns = ['position', 'word_original', 'word', 'first_position']
    df['is_new_word'] = df.apply(new_word, axis=1)
    df = df.reset_index(drop=True)
    # di = {1: "red", 0: "green"}
    # df['is_new_word'] = df['is_new_word'].map(di)
    original_words = df['word_original']
    colour_flag = df['is_new_word']

    document = Document()
    p = document.add_paragraph('')
    for idx, word in enumerate(original_words):
        if colour_flag[idx] == 1:
            p.add_run(' ' + original_words[idx]).bold = True
        else:
            p.add_run(' ' + original_words[idx])

    document.save('output.docx')
    return df


# -*- coding: utf-8 -*-
alphabets = "([A-Za-z])"
prefixes = "(Mr|St|Mrs|Ms|Dr)[.]"
suffixes = "(Inc|Ltd|Jr|Sr|Co)"
starters = "(Mr|Mrs|Ms|Dr|He\s|She\s|It\s|They\s|Their\s|Our\s|We\s|But\s|However\s|That\s|This\s|Wherever)"
acronyms = "([A-Z][.][A-Z][.](?:[A-Z][.])?)"
websites = "[.](com|net|org|io|gov)"


def split_into_sentences(text):
    text = " " + text + "  "
    text = text.replace("\n", " ")
    text = re.sub(prefixes, "\\1<prd>", text)
    text = re.sub(websites, "<prd>\\1", text)
    if "Ph.D" in text: text = text.replace("Ph.D.", "Ph<prd>D<prd>")
    text = re.sub("\s" + alphabets + "[.] ", " \\1<prd> ", text)
    text = re.sub(acronyms + " " + starters, "\\1<stop> \\2", text)
    text = re.sub(alphabets + "[.]" + alphabets + "[.]" + alphabets + "[.]", "\\1<prd>\\2<prd>\\3<prd>", text)
    text = re.sub(alphabets + "[.]" + alphabets + "[.]", "\\1<prd>\\2<prd>", text)
    text = re.sub(" " + suffixes + "[.] " + starters, " \\1<stop> \\2", text)
    text = re.sub(" " + suffixes + "[.]", " \\1<prd>", text)
    text = re.sub(" " + alphabets + "[.]", " \\1<prd>", text)
    if "”" in text: text = text.replace(".”", "”.")
    if "\"" in text: text = text.replace(".\"", "\".")
    if "!" in text: text = text.replace("!\"", "\"!")
    if "?" in text: text = text.replace("?\"", "\"?")
    text = text.replace(".", ".<stop>")
    text = text.replace("?", "?<stop>")
    text = text.replace("!", "!<stop>")
    text = text.replace("<prd>", ".")
    sentences = text.split("<stop>")
    sentences = sentences[:-1]
    sentences = [s.strip() for s in sentences]
    return sentences


def explode(df, lst_cols, fill_value='', preserve_index=False):
    # make sure `lst_cols` is list-alike
    if (lst_cols is not None
            and len(lst_cols) > 0
            and not isinstance(lst_cols, (list, tuple, np.ndarray, pd.Series))):
        lst_cols = [lst_cols]
    # all columns except `lst_cols`
    idx_cols = df.columns.difference(lst_cols)
    # calculate lengths of lists
    lens = df[lst_cols[0]].str.len()
    # preserve original index values
    idx = np.repeat(df.index.values, lens)
    # create "exploded" DF
    res = (pd.DataFrame({
        col: np.repeat(df[col].values, lens)
        for col in idx_cols},
        index=idx)
           .assign(**{col: np.concatenate(df.loc[lens > 0, col].values)
                      for col in lst_cols}))
    # append those rows that have empty lists
    if (lens == 0).any():
        # at least one list in cells is empty
        res = (res.append(df.loc[lens == 0, idx_cols], sort=False)
               .fillna(fill_value))
    # revert the original index order
    res = res.sort_index()
    # reset index if requested
    if not preserve_index:
        res = res.reset_index(drop=True)
    return res


def split_sentences(row):
    return row['sentence'].split()


def create_ordered_sentences_df(file_name):
    hp_text = file = open(file_name, encoding="utf-8-sig").read()
    sentences = split_into_sentences(hp_text)
    sentences_df = pd.DataFrame(sentences)
    sentences_df.columns = ['sentence']
    sentences_df['word'] = sentences_df.apply(split_sentences, axis=1)
    sentences_df = explode(sentences_df, 'word')
    sentences_df['word'] = sentences_df['word'].str.lower()
    sentences_df['word'] = sentences_df['word'].str.replace('[^\w\s]', '')
    word_freq_df = sentences_df.groupby('word').agg('count').reset_index().sort_values(by='sentence', ascending=False)
    word_freq_df = word_freq_df.reset_index(drop=True)
    word_freq_df = word_freq_df.drop(columns='sentence').reset_index()
    word_freq_df.columns = ['rank', 'word']
    sentences_df = sentences_df.merge(word_freq_df)
    sentences_df = sentences_df[['sentence', 'rank']].groupby('sentence').agg(max).reset_index()
    sentences_df = sentences_df.sort_values(by='rank').reset_index(drop=True)
    sentences_df['bare_sentence'] = sentences_df['sentence'].str.lower()
    sentences_df['bare_sentence'] = sentences_df['bare_sentence'].str.replace('[^\w\s]', '')
    sentences_df = sentences_df.reset_index()
    sentences_df.columns = ['bare_sentence_rank', 'sentence', 'sentence_difficulty', 'bare_sentence']
    bare_sentence_grouped_df = sentences_df[['bare_sentence', 'bare_sentence_rank']].groupby('bare_sentence').agg('min')
    bare_sentence_grouped_df = bare_sentence_grouped_df.reset_index()
    sentences_df = sentences_df.merge(bare_sentence_grouped_df)
    sentences_df = sentences_df[['sentence']]
    sentences_df.to_csv('easy_to_hard_sentences.csv')
    return sentences_df

def lookup_word_on_letonika(search_word):
    oh_dear = 1
    page_no = 0
    total_set = []
    sw_to_pf_list = []
    while oh_dear == 1:
        #print(search_word)
        url = 'https://www.letonika.lv/groups/default.aspx?r=1100&q=' + search_word + '&title=' + search_word + '/' + str(
            page_no) + '&g=5'
        res = requests.get(url)
        html_page = res.content
        soup = BeautifulSoup(html_page, 'html.parser')
        #text = soup.find_all(text=True)
        soup_txt = str(soup)
        if 'apstākļa vārds' in soup_txt:
            words = soup_txt.split('Entrytext">')
            if len(words) >= 4:
                words = words[1:4]
                for i in range(3):
                    words[i] = words[i].split('<')[0]
            else:
                words = []
        else:
            words = soup_txt.split('spelling')
            words = words[1::2]
            words = [w.replace('>', '').replace('<', '').replace('/', '') for w in words]
        if len(words) > 0:
            pamatforma = soup_txt.split('pamatforma: <i>')[1].split('<')[0]
            words.append(pamatforma)
            words = [pamatforma, words]
            total_set.append(words)
            sw_to_pf_list.append(pamatforma)
            page_no = page_no + 1
            print(search_word + ' - ' + pamatforma)
        else:
            oh_dear = 0
            # if idx % 100 == 99:
            #   total_list = list(total_set)
            #  df = pd.DataFrame(total_list)
            # df.to_csv('total_df.csv', index=False, encoding='utf-8')
    if len(sw_to_pf_list) == 0:
        print(search_word)
        total_set.append([search_word, [search_word]])
        sw_to_pf_list = [search_word]
    pf_to_conj_df = pd.DataFrame(total_set, columns=['pamatforma', 'conjugations']).explode('conjugations'). \
        drop_duplicates()
    sw_to_pf_df = pd.DataFrame()
    sw_to_pf_df['pamatforma'] = sw_to_pf_list
    sw_to_pf_df['search_word'] = search_word
    sw_to_pf_df = sw_to_pf_df.explode('pamatforma')
    return [pf_to_conj_df, sw_to_pf_df]

def lookup_multiple_words_on_letonika(word_list):

    dfs = [lookup_word_on_letonika(search_word) for search_word in word_list]
    pf_to_conjs_df = pd.concat([df[0] for df in dfs], ignore_index=True).drop_duplicates()
    sw_to_pf_df = pd.concat([df[1] for df in dfs], ignore_index=True)[['search_word','pamatforma']].drop_duplicates()

    return pf_to_conjs_df, sw_to_pf_df